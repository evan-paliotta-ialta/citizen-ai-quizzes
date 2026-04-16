#!/usr/bin/env python3
"""Generate the Citizen AI Engineer CIO Plan as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

output_path = "/Users/evanpaliotta/Desktop/iAltA Test/citizen ai engineer/CIO Plan - Citizen AI Engineer Program.docx"

doc = Document()

# Set narrow margins
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# Style defaults
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(9)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

DARK = RGBColor(0x1a, 0x1a, 0x2e)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
LIGHT_RED = RGBColor(0xFF, 0xF0, 0xF0)


def add_heading(text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK
        run.font.size = Pt(14 if level == 1 else 11)
    h.paragraph_format.space_before = Pt(2 if level == 1 else 8)
    h.paragraph_format.space_after = Pt(3)
    return h


def add_body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if bold_parts is None:
        run = p.add_run(text)
        run.font.size = Pt(9)
    else:
        # Parse simple bold markers **text**
        import re
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(9)
    return p


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def format_table(table, header_bg="1a1a2e", alt_rows=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_bg)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(8.5)
    # Data rows
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            if alt_rows and i % 2 == 1:
                set_cell_shading(cell, "F5F5F5")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)


def set_col_widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)


# === CONTENT ===

# Title
title = doc.add_heading("Citizen AI Engineer Program: Executive Plan", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# What We're Doing
add_heading("What We're Doing")
add_body(
    "Rolling out Claude AI licenses to employees using a **driver's license model**: "
    "everyone who receives a license completes a certification course, signs an acceptable use "
    "agreement, and follows tiered processes based on the sensitivity of their work. This enables "
    "productivity gains from AI while maintaining **SOC2 Type II** and **ISO 27001** compliance.",
    bold_parts=True
)

# Why Enterprise Plan is Required
add_heading("Why Enterprise Plan is Required")
add_body(
    "Our current Team plan provides technical guardrails (sandbox, managed settings, MCP server "
    "lockdown) but lacks the audit and monitoring capabilities our compliance certifications require. "
    "Enterprise adds:",
    bold_parts=True
)

table1 = doc.add_table(rows=5, cols=2, style="Table Grid")
data1 = [
    ("Capability", "Compliance Requirement"),
    ("Audit Logs", "SOC2 CC7.2 \u2014 Must prove who did what and when"),
    ("Compliance API (SIEM)", "SOC2 CC7.1 \u2014 Continuous monitoring required"),
    ("SCIM Provisioning", "SOC2 CC6.1/6.2 \u2014 Automated de-provisioning on offboard"),
    ("Custom Data Retention", "ISO 27001 A.8.10 \u2014 Provable data lifecycle management"),
]
for i, (c1, c2) in enumerate(data1):
    table1.rows[i].cells[0].text = c1
    table1.rows[i].cells[1].text = c2
set_col_widths(table1, [1.8, 4.8])
format_table(table1)

add_body(
    "**Without these detective controls, we can set rules but cannot prove they're being followed.**",
    bold_parts=True
)

# Driving Zones Model
add_heading("The Driving Zones Model")
add_body(
    "Usage is tiered by **data sensitivity** with no gray areas. The data determines the zone.",
    bold_parts=True
)

table2 = doc.add_table(rows=5, cols=3, style="Table Grid")
data2 = [
    ("Zone", "Data Involved", "Process Required"),
    ("Zone 1: Private Property\n(Your driveway)",
     "Public data only.\nNo company info.",
     "None \u2014 just use Claude."),
    ("Zone 2: Residential\n(Neighborhood streets)",
     "Internal, non-confidential.\nNo PII or client data.",
     "GitHub repo + branch protection\n+ code review."),
    ("Zone 3: Highway\n(High speed, high stakes)",
     "Confidential / PII / client data,\nMCP servers, production systems.",
     "Zone 2 + Jira ticket + approved\nMCP servers + full audit logging\n+ manager review."),
    ("No-Drive Zone\n(Zero tolerance)",
     "Credentials, secrets, bypassing\nguardrails, unapproved tools.",
     "Prohibited always.\nViolation = license revocation."),
]
for i, (c1, c2, c3) in enumerate(data2):
    table2.rows[i].cells[0].text = c1
    table2.rows[i].cells[1].text = c2
    table2.rows[i].cells[2].text = c3
set_col_widths(table2, [1.7, 2.4, 2.5])
format_table(table2)
# Red tint on No-Drive Zone row
for cell in table2.rows[4].cells:
    set_cell_shading(cell, "FFF0F0")

add_body(
    "**Override rule:** Any MCP server connecting to a live system automatically triggers Zone 3.",
    bold_parts=True
)

# Technical Guardrail Stack
add_heading("Technical Guardrail Stack")
add_body(
    "All enforced centrally via IT-deployed configuration files that users cannot override: "
    "**managed-settings.json** (org-wide permission rules, sandbox enforcement) \u00b7 "
    "**managed-mcp.json** (deny-by-default MCP server allowlist) \u00b7 "
    "**Global CLAUDE.md** (governance rules in every session) \u00b7 "
    "**Hooks** (audit logging, dangerous command blocking) \u00b7 "
    "**GitHub-first workflow** (all work in version control with branch protection) \u00b7 "
    "**Sandbox** (OS-level filesystem/network isolation \u2014 the true enforcement layer).",
    bold_parts=True
)

# Order of Operations
add_heading("Order of Operations")
table3 = doc.add_table(rows=6, cols=2, style="Table Grid")
data3 = [
    ("Phase", "Action"),
    ("1. Agree on the Plan", "CIO reviews against existing SOC2/ISO policies. Approve Driving Zones model, Enterprise upgrade, and overall approach."),
    ("2. Amend the Policies", "Update AUP to reflect Driving Zones, data classification rules, No-Drive violations, and enforcement procedures."),
    ("3. Build Technical Guardrails", "Deploy managed configs, upgrade to Enterprise, set up SIEM integration via Compliance API."),
    ("4. Design the Course", "Build certification: AI basics, Driving Zones, hands-on sandbox practice per zone, and assessment."),
    ("5. Pilot & Rollout", "5-10 person pilot across roles. Iterate. Then department-by-department rollout."),
]
for i, (c1, c2) in enumerate(data3):
    table3.rows[i].cells[0].text = c1
    table3.rows[i].cells[1].text = c2
    if i > 0:
        for p in table3.rows[i].cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
set_col_widths(table3, [1.6, 5.0])
format_table(table3)

# The Ask
add_heading("The Ask")
ask_text = (
    "1. Approve the Driving Zones tiering model as our governance framework\n"
    "2. Approve upgrade from Claude Team to Claude Enterprise plan\n"
    "3. Review the Acceptable Use Policy draft against this framework\n"
    "4. Designate compliance/legal reviewers to align with SOC2/ISO control matrices"
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
run = p.add_run(ask_text)
run.font.size = Pt(9)

doc.save(output_path)
print(f"Word doc generated: {output_path}")
