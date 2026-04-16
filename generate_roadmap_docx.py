#!/usr/bin/env python3
"""Generate the Citizen AI Engineer Roadmap as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

output_path = "/Users/evanpaliotta/Desktop/iAltA Test/Citizen AI Engineer/Citizen AI Engineer Program - Roadmap.docx"

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Style defaults
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(9.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

DARK      = RGBColor(0x1a, 0x1a, 0x2e)
BLUE      = RGBColor(0x2F, 0x54, 0x96)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xF5, 0xF7, 0xFA)
YELLOW_BG = RGBColor(0xFF, 0xFB, 0xE6)
RED_BG    = RGBColor(0xFF, 0xF0, 0xF0)


def add_heading(text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK
        run.font.name = "Calibri"
        run.font.size = Pt(13 if level == 1 else 11)
        run.font.bold = True
    h.paragraph_format.space_before = Pt(10 if level == 2 else 4)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_body(text, bold_parts=True):
    import re
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
    return p


def add_bullet(text, bold_parts=True, indent=1):
    import re
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25 * indent)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
    return p


def set_cell_shading(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    tcPr.append(shd)


def format_table(table, header_bg="1a1a2e"):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            if i % 2 == 1:
                set_cell_shading(cell, "F5F7FA")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"


def set_col_widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)


def add_phase_label(phase_num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"PHASE {phase_num}  —  {title.upper()}")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = BLUE
    return p


def divider():
    p = doc.add_paragraph("─" * 95)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(7)


# ═══════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════
title = doc.add_heading("Citizen AI Engineer Program", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = DARK
    run.font.size = Pt(16)
    run.font.name = "Calibri"

subtitle = doc.add_paragraph("Implementation Roadmap")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(4)
for run in subtitle.runs:
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.italic = True

divider()

# ═══════════════════════════════════════════════════════
# PROGRAM OVERVIEW
# ═══════════════════════════════════════════════════════
add_heading("Program Overview")
add_body(
    "The Citizen AI Engineer initiative empowers every non-technical employee at iAltA to use Claude AI "
    "safely, productively, and in full compliance with our SOC2 Type II and ISO 27001 certifications. "
    "Employees receive a **single AI license** upon completing a comprehensive certification program and signing "
    "the Acceptable Use Policy. The license grants full access to Claude, including MCP servers and agentic "
    "workflows. Governance is handled by the **Driving Zones framework**, which categorizes work by data "
    "sensitivity and enforces appropriate process requirements at the system level — not by restricting "
    "which tools employees can access."
)

add_heading("The Driving Zones Model", level=2)
add_body("Every task falls into one of four zones determined by the data involved, not the tool used.")

zones_table = doc.add_table(rows=5, cols=3, style="Table Grid")
zones_data = [
    ("Zone", "Data Involved", "Process Required"),
    ("Zone 1 — Private Property", "Public data only. No company information.", "None — just use Claude."),
    ("Zone 2 — Residential", "Internal, non-confidential. No PII or client data.", "GitHub repo + branch protection + code review."),
    ("Zone 3 — Highway", "Confidential, PII, client data, MCP servers, production systems.", "Zone 2 requirements + Jira ticket + approved MCP servers + audit logging + manager review."),
    ("No-Drive Zone", "Credentials, secrets, bypassing guardrails, unapproved tools.", "Prohibited always. Violation = license revocation."),
]
for i, (c1, c2, c3) in enumerate(zones_data):
    zones_table.rows[i].cells[0].text = c1
    zones_table.rows[i].cells[1].text = c2
    zones_table.rows[i].cells[2].text = c3
set_col_widths(zones_table, [1.6, 2.5, 2.6])
format_table(zones_table)
for cell in zones_table.rows[4].cells:
    set_cell_shading(cell, "FFF0F0")

add_body("**Override rule:** Any MCP server connecting to a live system automatically triggers Zone 3, regardless of the data involved.")

divider()

# ═══════════════════════════════════════════════════════
# PHASE 0
# ═══════════════════════════════════════════════════════
add_phase_label(0, "Get the Foundation Approved")
add_body(
    "All subsequent phases depend on the decisions made here. Nothing is built until the plan is approved "
    "and the infrastructure decision is made."
)
for task in [
    "CIO decision on **Enterprise vs Teams** plan",
    "Enterprise procurement (minimum 20 seats; self-serve at anthropic.com or sales-assisted)",
    "Legal/InfoSec drafts the **Acceptable Use Policy (AUP)**",
    "Update the Information Security Policy to include AI tools",
    "Brief Board of Directors",
]:
    add_bullet(task)

divider()

# ═══════════════════════════════════════════════════════
# PHASE 1
# ═══════════════════════════════════════════════════════
add_phase_label(1, "Build the Guardrails")
add_body(
    "Technical controls are deployed centrally and cannot be overridden by individual users. "
    "These must be in place before any employee accesses Claude."
)
for task in [
    "Deploy **managed-settings.json** + **managed-mcp.json** via enterprise admin console",
    "Write company-wide **CLAUDE.md** template (Jira-first, no credentials, Zone rules baked in)",
    "Configure **PreToolUse blocking hooks** (block: rm -rf, push to main, secret/credential patterns)",
    "Configure **PostToolUse audit hooks** (log all tool use to SIEM via Compliance API)",
    "Set up **SCIM provisioning** — directly remediates SOC2 audit exceptions CC5.1.1 and CC6.1.2",
    "Verify audit log pipeline is flowing correctly to security tooling",
]:
    add_bullet(task)

divider()

# ═══════════════════════════════════════════════════════
# PHASE 2
# ═══════════════════════════════════════════════════════
add_phase_label(2, "Build the Certification Program")
add_body(
    "The certification program is the gateway to Claude access. Every employee completes the full "
    "curriculum — all three modules — before receiving a license. There are no partial licenses or tiered "
    "access levels. Think of it as a driver's license: you either have it or you don't, but earning it "
    "requires demonstrating knowledge of all roads."
)

add_heading("Course Strategy Decision", level=2)
add_body("Before building, the program lead must decide on the delivery model:")

strategy_table = doc.add_table(rows=4, cols=3, style="Table Grid")
strategy_data = [
    ("Option", "Approach", "Trade-offs"),
    (
        "A — Anthropic Academy Only",
        "Use free Anthropic Academy courses as the curriculum backbone. Employees submit proof of completion.",
        "Low build cost. Dependent on Anthropic's catalog staying current and available.",
    ),
    (
        "B — Internal Only",
        "Build all course content in-house using Claude. Full control over content and tone.",
        "Higher upfront effort. You own and maintain everything indefinitely.",
    ),
    (
        "C — Hybrid (Recommended)",
        "Anthropic Academy covers AI fundamentals. Internal courses cover iAltA-specific rules, Driving Zones, workflows, and hands-on exercises.",
        "Best balance of build effort and content relevance. Lower dependency on a single source.",
    ),
]
for i, (c1, c2, c3) in enumerate(strategy_data):
    strategy_table.rows[i].cells[0].text = c1
    strategy_table.rows[i].cells[1].text = c2
    strategy_table.rows[i].cells[2].text = c3
set_col_widths(strategy_table, [1.6, 2.6, 2.5])
format_table(strategy_table)
for cell in strategy_table.rows[3].cells:
    set_cell_shading(cell, "F0F7E6")

add_heading("Proof of Completion Process", level=2)
add_body("If using Anthropic Academy courses, the following process applies:")
for step in [
    "Employee completes designated Academy course(s) and downloads their completion certificate",
    "Employee submits certificate to the central tracker (HubSpot form, Google Sheet, or equivalent) with their name and level",
    "Program lead reviews submission and marks the employee as certified in the tracker",
    "Claude access is provisioned via SCIM only after the certification record is confirmed",
]:
    add_bullet(step)

add_heading("One Certification. One License.", level=2)
add_body(
    "All three modules are required. Completion of the full curriculum, passing the assessment, "
    "and signing the AUP results in one license granting full Claude access."
)

modules_table = doc.add_table(rows=4, cols=3, style="Table Grid")
modules_data = [
    ("Module", "Content", "Delivery"),
    (
        "Module 1\nAI Fundamentals",
        "How LLMs work. What AI can and can't do. Prompt basics. Setting expectations for AI-assisted work.",
        "Anthropic Academy course + short quiz.",
    ),
    (
        "Module 2\nResponsible Use at iAltA",
        "Driving Zones in full (Zones 1, 2, 3, No-Drive). The AUP. Data classification. Real iAltA examples and edge cases.",
        "Internal course + quiz. AUP signature required.",
    ),
    (
        "Module 3\nAdvanced Workflows",
        "MCP servers: what they are, which are approved, why they auto-trigger Zone 3. Agentic task design. Jira-first enforcement. CLAUDE.md awareness. Zone 3 in practice.",
        "Internal course + practical exercise.",
    ),
]
for i, (c1, c2, c3) in enumerate(modules_data):
    modules_table.rows[i].cells[0].text = c1
    modules_table.rows[i].cells[1].text = c2
    modules_table.rows[i].cells[2].text = c3
set_col_widths(modules_table, [1.3, 3.3, 2.1])
format_table(modules_table)

add_heading("Curriculum Build Tasks", level=2)
for task in [
    "Make the course strategy decision (Option A, B, or C above)",
    "Audit the current **Anthropic Academy catalog** and designate specific courses for Module 1 (and Module 2 if applicable)",
    "Build **Module 2** internal course: iAltA Driving Zones, AUP walkthrough, data classification scenarios",
    "Build **Module 3** internal course: MCP servers, agentic workflows, Jira ticket process, hands-on exercise",
    "Build assessments and quizzes for each module",
    "Create the **completion tracker** and proof-of-completion intake process",
    "Design the **license record** — a signed, dated completion document per employee",
]:
    add_bullet(task)

divider()

# ═══════════════════════════════════════════════════════
# PHASE 3
# ═══════════════════════════════════════════════════════
add_phase_label(3, "Company-Wide Rollout")
add_body(
    "Once guardrails are live and the certification program is built, the program goes live for all employees "
    "simultaneously. Using AI without a license is a policy violation from day one."
)
for task in [
    "All employees complete the full certification before Claude access is provisioned — no exceptions",
    "SCIM-based provisioning configured and tested: access is tied to the certification record, no manual onboarding",
    "Company-wide announcement via all-hands: program is now live, policy is in effect",
    "Central certification tracker live and being maintained",
]:
    add_bullet(task)

divider()

# ═══════════════════════════════════════════════════════
# PHASE 4
# ═══════════════════════════════════════════════════════
add_phase_label(4, "Steady State")
add_body(
    "The following annual touchpoints apply regardless of which ongoing engagement model is selected:"
)
for task in [
    "**Annual re-certification** — all employees re-read the AUP and re-sign",
    "**Annual vendor compliance review** — obtain and review updated Anthropic SOC2 Type II report",
    "**Annual curriculum review** — update course content for new Claude capabilities, policy changes, and new zones",
    "**Annual risk assessment update**",
    "**Incident response drill** — run a tabletop exercise with an AI-specific scenario",
]:
    add_bullet(task)

add_heading("Ongoing Engagement Model — CIO to Select", level=2)
add_body("Choose one of the following models for how the program lead engages with employees on an ongoing basis:")

# Option 1
add_heading("Option 1 — 1-on-1 Quarterly Usage Reviews", level=2)
add_body(
    "The program lead meets individually with each licensed employee once per quarter for a 1-hour session. "
    "Suggested structure:"
)
for item in [
    "**15 min** — Review actual usage from audit logs: tools used, zones invoked, frequency",
    "**20 min** — Open Q&A: what's working, what's confusing, where they're avoiding AI because they're unsure of the rules",
    "**15 min** — Coaching: suggest better workflows, correct Zone misclassifications, identify employees ready for more advanced use",
    "**10 min** — Flag and log any policy edge cases for future curriculum updates",
]:
    add_bullet(item)
add_body(
    "**Best for:** Surfacing real usage patterns and catching misuse early. Builds individual trust and adoption. "
    "Time cost scales linearly with headcount — plan accordingly."
)

# Option 2
add_heading("Option 2 — Weekly Office Hours", level=2)
add_body(
    "The program lead holds a standing 1-hour open session each week. No agenda, no required attendance — "
    "employees drop in with questions, use cases they're unsure about, or things they want to try but "
    "aren't sure are permitted. Pairs well with a shared Slack/Teams channel for async questions between "
    "sessions and a running FAQ log built from recurring questions."
)
add_body(
    "**Best for:** Fast feedback loop during early adoption with minimal scheduling overhead. "
    "Requires consistent presence but no prep work."
)

# Option 3
add_heading("Option 3 — Monthly Mandatory Group Sessions", level=2)
add_body("All employees attend a 1-hour monthly session. Suggested structure:")
for item in [
    "**15 min** — Usage recap: aggregate trends from audit logs (anonymized), common patterns, wins",
    "**20 min** — Case study or demo: one real use case from the past month, instructive or cautionary",
    "**15 min** — Policy/curriculum update: anything that changed, newly approved MCP servers, zone clarifications",
    "**10 min** — Open Q&A",
]:
    add_bullet(item)
add_body(
    "**Best for:** Keeping the whole company aligned and building a shared AI culture. "
    "Efficient as headcount grows. Less personalized than 1-on-1s — individual edge cases may stay hidden."
)

# Option 4
add_heading("Option 4 — Tiered Engagement by Usage Level", level=2)
add_body("Match engagement intensity to how actively and riskily each employee uses Claude:")
for item in [
    "**Light users** — monthly group session only (low risk, broad audience)",
    "**Regular users** — monthly group + quarterly 30-minute 1-on-1",
    "**Heavy / agentic users** — monthly group + monthly 30-minute 1-on-1 + ad-hoc audit log reviews",
]:
    add_bullet(item)
add_body(
    "**Best for:** Focusing program lead time where complexity and risk are highest while keeping everyone informed."
)

# Option 5
add_heading("Option 5 — Async-First with Periodic Escalation", level=2)
add_body(
    "The primary channel is a dedicated Slack/Teams channel monitored by the program lead. "
    "A monthly written digest covers top questions, notable usage patterns, policy reminders, and curriculum updates. "
    "Live sessions are reserved for policy incidents, major Claude capability changes, or annual re-certification."
)
add_body(
    "**Best for:** Lean organizations where meeting overhead is a concern. "
    "Lower signal on adoption issues — problems may surface later than with other models."
)

divider()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run = p.add_run(
    "Recommendation: Option 4 (Tiered by Usage) balances program lead time against risk proportionality. "
    "Option 2 (Office Hours) pairs well with any choice as a low-cost safety valve during the first year."
)
run.font.size = Pt(9)
run.font.italic = True
run.font.name = "Calibri"
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.save(output_path)
print(f"Roadmap Word doc generated: {output_path}")
