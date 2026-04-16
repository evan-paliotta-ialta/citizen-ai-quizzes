#!/usr/bin/env python3
"""Generate: Citizen AI Engineer Program - Launch Plan.docx"""

import os, io
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
OUT  = os.path.join(BASE, "Citizen AI Engineer Program - Launch Plan.docx")

# ─── Colours ──────────────────────────────────────────────────────────────────
DARK  = RGBColor(0x0F, 0x11, 0x17)
BLUE  = RGBColor(0x2D, 0x6F, 0xF7)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY = RGBColor(0x44, 0x44, 0x44)
GREEN = RGBColor(0x1A, 0x7A, 0x3E)
GOLD  = RGBColor(0xB8, 0x76, 0x00)
RED   = RGBColor(0xC0, 0x39, 0x2B)
HBLU  = RGBColor(0x1A, 0x48, 0xBB)
PURP  = RGBColor(0x7B, 0x3F, 0xB5)

# ══════════════════════════════════════════════════════════════════════════════
# ARCHITECTURE DIAGRAM
# CIO at top, citizens at bottom, white background
# ══════════════════════════════════════════════════════════════════════════════
def build_diagram():
    fig, ax = plt.subplots(figsize=(15, 10))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 10)
    ax.set_facecolor("#FFFFFF")
    fig.patch.set_facecolor("#FFFFFF")
    ax.axis("off")

    C_BLUE    = "#2D6FF7"
    C_BLUE_DK = "#1A48BB"
    C_GREEN   = "#1A7A3E"
    C_PURPLE  = "#7B3FB5"
    C_AMBER   = "#C07800"
    C_RED     = "#C0392B"
    C_WHITE   = "#FFFFFF"
    C_LGRAY   = "#555555"
    C_MGRAY   = "#888888"
    C_PANEL   = "#F4F6FB"
    C_PANEL2  = "#EEF2FF"

    def box(x, y, w, h, fc, ec=None, lw=1.2, ls="-", radius=0.25, alpha=1.0):
        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle=f"round,pad=0,rounding_size={radius}",
                              facecolor=fc, edgecolor=ec or fc,
                              linewidth=lw, linestyle=ls, alpha=alpha)
        ax.add_patch(rect)

    def outline(x, y, w, h, ec, lw=1.8, ls="--", radius=0.3):
        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle=f"round,pad=0,rounding_size={radius}",
                              facecolor="none", edgecolor=ec,
                              linewidth=lw, linestyle=ls)
        ax.add_patch(rect)

    def t(x, y, s, size=9, color=C_LGRAY, bold=False, ha="center", va="center", italic=False):
        weight = "bold" if bold else "normal"
        style  = "italic" if italic else "normal"
        ax.text(x, y, s, fontsize=size, color=color, ha=ha, va=va,
                fontweight=weight, fontstyle=style, fontfamily="DejaVu Sans")

    def arrow(x1, y1, x2, y2, color=C_BLUE, lw=1.6):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=color, lw=lw,
                                   mutation_scale=14,
                                   connectionstyle="arc3,rad=0.0"))

    def label_arrow(x1, y1, x2, y2, label, color=C_MGRAY, lw=1.4, offset=(0.12, 0)):
        arrow(x1, y1, x2, y2, color=color, lw=lw)
        mx = (x1+x2)/2 + offset[0]
        my = (y1+y2)/2 + offset[1]
        t(mx, my, label, size=7.5, color=color, italic=True, ha="left")

    # ── ROW 1 — CIO DASHBOARD  (top, y=8.7–9.7) ─────────────────────────────
    box(2.8, 8.7, 9.4, 0.9, C_BLUE_DK, radius=0.3)
    t(7.5, 9.19, "CIO Dashboard  ·  Board Metrics", 12, C_WHITE, bold=True)
    t(7.5, 8.87, '"X citizens shipped Y tools → closed Z bottlenecks → N% reduction in support load"',
      8, "#BDD3FF")

    # ── ROW 2 — HIGHLANDER  (y=6.35–8.45) ───────────────────────────────────
    box(0.3, 6.35, 14.4, 2.1, C_PANEL, ec="#C0A0E8", lw=1.2, radius=0.35)
    t(7.5, 8.28, "Highlander ODL  —  Ops Data Lake", 11, C_PURPLE, bold=True)

    sub_boxes = [
        (0.55, 6.5,  4.3, "Ingestion",
         "GitHub webhook → staging.github_commits\nS3 raw data · Aurora PostgreSQL"),
        (5.1,  6.5,  4.6, "Analyzers",
         "engineer_effectiveness · gap_analysis\nOKR attribution · constraint analysis"),
        (10.0, 6.5,  4.5, "Metrics Output",
         "Commit velocity · AI co-author %\nCert funnel · OKR attribution"),
    ]
    for bx, by, bw, title, sub in sub_boxes:
        box(bx, by, bw, 1.55, C_WHITE, ec="#C0A0E8", lw=1.0, radius=0.2)
        t(bx+bw/2, by+1.2,  title, 9,   C_PURPLE, bold=True)
        t(bx+bw/2, by+0.62, sub,   7.5, C_LGRAY)

    label_arrow(7.5, 8.68, 7.5, 8.47, "  real-time metrics", color=C_PURPLE, lw=1.8)

    # ── ROW 3 — GITHUB ORG  (y=4.35–6.1) ────────────────────────────────────
    box(1.4, 4.35, 12.2, 1.75, C_PANEL, ec=C_GREEN, lw=1.5, radius=0.3)
    t(7.5, 5.84, "GitHub Teams Org  —  Citizen Repositories", 11, C_GREEN, bold=True)
    t(7.5, 5.52, "citizen-<team>-<project>  ·  METADATA.yaml links every repo to an OKR  ·  branch protection",
      8.5, C_LGRAY)
    t(7.5, 5.24, "Repo owner controls merges via branch protection  ·  PR review required before any merge",
      8,   C_MGRAY)
    t(7.5, 4.62, "90-day org-level audit log  ·  SOC2 compensating control  ·  rollback for any AI mistake",
      7.5, C_MGRAY, italic=True)

    label_arrow(7.5, 6.33, 7.5, 6.13, "  webhook (real-time)", color=C_GREEN, lw=1.8)

    # ── ROW 4 — GUARDRAIL WRAPPER  (y=1.55–4.12) ─────────────────────────────
    outline(0.2, 1.55, 14.6, 2.57, C_AMBER, lw=2.2, ls="--", radius=0.4)
    t(7.5, 4.28, "GUARDRAIL STACK", 9.5, C_AMBER, bold=True)

    # Guardrail pills
    pills = [
        "managed-settings.json",
        "managed-mcp.json",
        "Global CLAUDE.md",
        "PreToolUse hooks",
        "Sandbox",
    ]
    pw, ph = 2.7, 0.33
    pg = 0.23
    total = len(pills)*pw + (len(pills)-1)*pg
    px0   = (15 - total) / 2
    for i, pill in enumerate(pills):
        px = px0 + i*(pw+pg)
        box(px, 3.68, pw, ph, "#FFF3CD", ec=C_AMBER, lw=0.9, radius=0.12)
        t(px+pw/2, 3.845, pill, 7.5, C_AMBER)

    # Claude Desktop
    box(0.45, 1.7, 5.6, 1.82, C_PANEL2, ec=C_BLUE, lw=1.3, radius=0.25)
    box(0.45, 3.12, 5.6, 0.4, C_BLUE, ec=C_BLUE, lw=0, radius=0.2)
    t(3.25, 3.32, "Claude Desktop", 11, C_WHITE, bold=True)
    t(3.25, 2.92, "GUI — no terminal required", 8.5, C_LGRAY)
    t(3.25, 2.65, "Non-technical citizens", 8,   C_LGRAY)
    t(3.25, 2.38, "Visual diffs · Cowork background agents", 7.5, C_MGRAY)
    t(3.25, 2.1,  "One-click MCP Desktop Extensions (.mcpb)", 7.5, C_MGRAY)
    t(3.25, 1.85, "Recommended for most citizens", 7.5, C_BLUE, italic=True)

    # Claude Code
    box(6.2,  1.7, 5.6, 1.82, C_PANEL2, ec=C_BLUE_DK, lw=1.3, radius=0.25)
    box(6.2,  3.12, 5.6, 0.4, C_BLUE_DK, ec=C_BLUE_DK, lw=0, radius=0.2)
    t(9.0,  3.32, "Claude Code", 11, C_WHITE, bold=True)
    t(9.0,  2.92, "Terminal / CLI", 8.5, C_LGRAY)
    t(9.0,  2.65, "More technical citizens", 8,   C_LGRAY)
    t(9.0,  2.38, "Full Git automation · managed-settings stack", 7.5, C_MGRAY)
    t(9.0,  2.1,  "Branch, commit, PR without GitHub UI", 7.5, C_MGRAY)
    t(9.0,  1.85, "Strongest compliance posture", 7.5, C_BLUE_DK, italic=True)

    # MCP servers panel
    box(12.0, 1.7, 2.65, 1.82, C_PANEL, ec=C_BLUE, lw=0.8, radius=0.2)
    t(13.33, 3.3, "MCP", 8.5, C_BLUE, bold=True)
    t(13.33, 3.08, "Servers", 8.5, C_BLUE, bold=True)
    t(13.33, 2.78, "26 official", 7.5, C_LGRAY)
    t(13.33, 2.57, "connectors", 7.5, C_LGRAY)
    t(13.33, 2.3,  "Highlander", 7.5, C_LGRAY)
    t(13.33, 2.1,  "Playwright", 7.5, C_LGRAY)
    t(13.33, 1.88, "Others: written", 7, C_MGRAY, italic=True)
    t(13.33, 1.73, "approval req.", 7, C_MGRAY, italic=True)

    # MCP arrows
    ax.annotate("", xy=(11.98, 2.5), xytext=(6.06, 2.5),
                arrowprops=dict(arrowstyle="<->", color=C_BLUE, lw=1.0,
                                mutation_scale=10))
    ax.annotate("", xy=(11.98, 2.7), xytext=(6.06, 2.7),
                arrowprops=dict(arrowstyle="<->", color=C_BLUE_DK, lw=0.8,
                                mutation_scale=8))

    # arrows into GitHub from tools
    label_arrow(3.25, 4.33, 5.4,  4.2,  "  commits", color=C_GREEN, lw=1.5)
    label_arrow(9.0,  4.33, 8.6,  4.2,  "  commits", color=C_GREEN, lw=1.5)

    # ── ROW 5 — CITIZENS  (bottom, y=0.05–1.3) ───────────────────────────────
    citizen_data = [
        (1.1,  "Ops"),
        (3.85, "Finance"),
        (6.6,  "Sales"),
        (9.35, "Product"),
        (12.1, "Any team"),
    ]
    for cx, lbl in citizen_data:
        box(cx-0.55, 0.1, 1.1, 1.1, C_PANEL, ec="#CCCCCC", lw=0.8, radius=0.2)
        t(cx, 0.96, "●",  11,  C_LGRAY)
        t(cx, 0.67, "▲",  13,  C_LGRAY)
        t(cx, 0.25, lbl,  7.5, C_MGRAY)

    t(7.5, 1.43, "CITIZENS", 9.5, C_LGRAY, bold=True)

    # arrows citizens → guardrail zone
    for cx, _ in citizen_data:
        arrow(cx, 1.22, cx, 1.57, color="#AAAAAA", lw=1.0)

    # ── LEGEND ───────────────────────────────────────────────────────────────
    legend = [
        (C_BLUE,   "Claude AI tooling"),
        (C_AMBER,  "Guardrail stack"),
        (C_GREEN,  "GitHub Org"),
        (C_PURPLE, "Highlander ODL"),
    ]
    lx, ly = 0.22, 5.8
    for clr, lbl in legend:
        box(lx, ly-0.1, 0.22, 0.22, clr, radius=0.04)
        t(lx+0.34, ly+0.02, lbl, 7.5, C_LGRAY, ha="left")
        ly -= 0.36

    plt.tight_layout(pad=0)
    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=180, bbox_inches="tight",
                facecolor="#FFFFFF", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# WORD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin   = Inches(0.85)
    section.right_margin  = Inches(0.85)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after  = Pt(3)
style.paragraph_format.space_before = Pt(0)


def set_cell_bg(cell, hex_str):
    tcPr = cell._element.get_or_add_tcPr()
    shd  = tcPr.makeelement(qn("w:shd"),
                             {qn("w:fill"): hex_str, qn("w:val"): "clear"})
    tcPr.append(shd)


def cell_write(cell, text, size=9.5, bold=False, color=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.name   = "Calibri"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color or LGRAY


def col_widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text.upper())
    run.font.name  = "Calibri"
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = BLUE
    return p


def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(10)
    run.font.color.rgb = LGRAY
    return p


def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.2)
    run = p.add_run(text)
    run.font.name   = "Calibri"
    run.font.size   = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    return p


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 100)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def hyperlink_run(para, text, url, size=9.5, color=None):
    """Add a hyperlink run to an existing paragraph."""
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                           is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)

    if color:
        clr = OxmlElement("w:color")
        clr.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
        rPr.append(clr)

    new_run.append(rPr)
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    new_run.append(t_elem)
    hyperlink.append(new_run)
    para._p.append(hyperlink)


# ─── Title ────────────────────────────────────────────────────────────────────
title = doc.add_heading("Citizen AI Engineer Program", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in title.runs:
    run.font.color.rgb = DARK
    run.font.size      = Pt(18)
    run.font.name      = "Calibri"

sub = doc.add_paragraph("Launch Plan  ·  April 2026  ·  Confidential")
for run in sub.runs:
    run.font.name      = "Calibri"
    run.font.size      = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.italic    = True

divider()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(10)
run = p.add_run(
    "We are rolling out Claude AI to every employee who wants it using a driver's license model: "
    "learn the rules, sign the agreement, get the license. "
    "The goal is not to restrict — it is to give people the tools to offload tasks to AI "
    "so they can focus on the actual purpose of their role. "
    "Most of the program is already built. What remains is a handful of procurement, approval, and build items."
)
run.font.name      = "Calibri"
run.font.size      = Pt(10.5)
run.font.color.rgb = LGRAY

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT'S DONE
# ══════════════════════════════════════════════════════════════════════════════
h1("What's Already Done")

done_tbl = doc.add_table(rows=6, cols=2)
done_tbl.style = "Table Grid"
col_widths(done_tbl, [3.0, 4.77])

done_rows = [
    ("Item", "Detail"),
    ("Claude Teams plan",  "Contracted. 25 seats available immediately."),
    ("Course content",     "Bulk of curriculum built. Hosted on SharePoint (ialta.sharepoint.com/sites/CitizenAI). Remaining modules in progress."),
    ("Program design",     "Driving Zones, AUP draft, MCP allowlist, SOC2 control map — all documented."),
    ("Tracking engine",    "Highlander ODL ingests GitHub commits in real time. Connecting citizens requires a one-time setup."),
    ("Program lead",       "Designated. Handles course, office hours, 1:1s, and ongoing management."),
]
for i, (c1, c2) in enumerate(done_rows):
    row = done_tbl.rows[i]
    if i == 0:
        set_cell_bg(row.cells[0], "2D6FF7")
        set_cell_bg(row.cells[1], "2D6FF7")
        cell_write(row.cells[0], c1, bold=True, color=WHITE)
        cell_write(row.cells[1], c2, bold=True, color=WHITE)
    else:
        fill = "F0F4FF" if i % 2 == 1 else "FFFFFF"
        set_cell_bg(row.cells[0], fill)
        set_cell_bg(row.cells[1], fill)
        cell_write(row.cells[0], "✓  " + c1, bold=True, color=GREEN)
        cell_write(row.cells[1], c2)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — STEP BY STEP
# ══════════════════════════════════════════════════════════════════════════════
h1("Step-by-Step Launch Plan")

body(
    "Steps are ordered by dependency. Items marked LEADERSHIP are where your involvement directly unblocks progress."
)

# We'll write the table, then add pricing note paragraphs inline for step 1.
steps = [
    (
        "1", "LEADERSHIP",
        "Approve GitHub Teams procurement",
        (
            "GitHub Teams is the version control and audit trail infrastructure for the program. "
            "Every citizen's AI-assisted work lives here — it feeds Highlander and provides the SOC2 compensating control.\n\n"
            "GitHub Teams: $4.00/user/month (monthly billing) or $3.67/user/month (annual billing).\n"
            "Source: github.com/pricing\n\n"
            "Claude Teams (already contracted): $25/seat/month (standard) or $125/seat/month (premium — includes Claude Code).\n"
            "Source: claude.com/pricing"
        ),
        "1A48BB", HBLU,
    ),
    (
        "2", "LEGAL / IT LEADERSHIP",
        "Get the Acceptable Use Policy reviewed and signed off",
        (
            "The AUP draft is complete. It defines the Driving Zones model and governs all employee AI use. "
            "Needs legal and IT leadership review before it can be issued to employees. "
            "This gates license issuance — no employee officially receives access until the AUP is signed off."
        ),
        "B87600", GOLD,
    ),
    (
        "3", "PROGRAM LEAD",
        "Finish the course",
        (
            "Remaining modules need to be completed and published on SharePoint (ialta.sharepoint.com/sites/CitizenAI). "
            "Priority modules: Operating Framework (Driving Zones in full), GitHub for non-technical users, "
            "and Safety & Responsible Use.\n\n"
            "The Jensen Huang mental model — a job is a purpose plus a set of tasks; "
            "AI offloads tasks so you can find more problems to solve — should be woven into the opening module "
            "as the core philosophy of the program. The course must be complete before the company-wide announcement."
        ),
        "1A7A3E", GREEN,
    ),
    (
        "4", "PROGRAM LEAD",
        "Deploy the guardrail stack",
        (
            "See the Guardrails section below for the full specification. "
            "All controls are confirmed feasible on the Claude Teams plan — no Enterprise upgrade required.\n\n"
            "Two deployment paths:\n"
            "1. Admin console (Admin Settings > Claude Code > Managed Settings): pushes permission rules, "
            "hooks, and sandbox config to all authenticated users. This is a public beta feature, available on Teams v2.1.38+.\n"
            "2. File-based (stable, works on any plan): IT drops managed-settings.json and managed-mcp.json "
            "into system directories on each machine. Both Claude Code (terminal) and Claude Desktop read "
            "from the same files — one deployment covers both.\n\n"
            "Important: MCP server control (managed-mcp.json) cannot be pushed through the admin console — "
            "it must be deployed via the file-based approach. For 25 people, a simple onboarding script "
            "that citizens run once during setup is the right path. The script pulls both files from a "
            "private GitHub repo and drops them in the correct system directory for the citizen's OS.\n\n"
            "CLAUDE.md is committed to every citizen repo template and auto-loads when they open that folder — "
            "no deployment needed."
        ),
        "555555", LGRAY,
    ),
    (
        "5", "PROGRAM LEAD",
        "Stand up the GitHub org for citizens",
        (
            "Create a dedicated GitHub Teams org, separate from the engineering org. "
            "Build the citizen repo template including METADATA.yaml (links each repo to an OKR), "
            "a README structure, and branch protection rules.\n\n"
            "Branch protection gives each repo owner full control over what gets merged — "
            "PRs require owner review before any change is accepted. "
            "This is the foundation for collaboration as citizens grow more comfortable: "
            "citizens can open PRs on one another's repos, and the owner decides what goes in.\n\n"
            "Naming convention: citizen-<team>-<project>. Topic tag: citizen-ai (used by Highlander to segment data).\n\n"
            "Note on Global CLAUDE.md: The project-level CLAUDE.md committed to every repo template "
            "is the recommended approach on the Teams plan — it auto-loads when a citizen opens that folder "
            "in Claude Code and requires zero additional setup. See the Guardrails section for detail."
        ),
        "555555", LGRAY,
    ),
    (
        "6", "CIO",
        "Connect citizens to Highlander",
        (
            "Register citizen repos in admin.github_repositories, map identities in engineer_email_mapping, "
            "and add a 'citizen' segment to the engineer_effectiveness analyzer. "
            "After this, every commit a citizen makes flows into Highlander automatically — "
            "no manual reporting. OKR attribution comes from the METADATA.yaml in each repo."
        ),
        "7B3FB5", PURP,
    ),
    (
        "7", "LEADERSHIP",
        "Company-wide announcement",
        (
            "Leadership endorsement is the single biggest adoption accelerator. "
            "Message: the program is live, the course is required for anyone using Claude, "
            "here is the link, 30 days for existing users to complete it. "
            "This brings currently ungoverned Claude users under the program umbrella."
        ),
        "1A48BB", HBLU,
    ),
    (
        "8", "PROGRAM LEAD",
        "Open enrolment and begin onboarding",
        (
            "Course live on SharePoint. Citizens complete the curriculum, sign the AUP, receive their license. "
            "Weekly open office hours begin immediately. "
            "Highlander starts tracking commit activity from the first push."
        ),
        "555555", LGRAY,
    ),
    (
        "9", "PROGRAM LEAD",
        "Run the ongoing cadence",
        (
            "Weekly open office hours (optional, 30 min).\n\n"
            "Monthly all-hands call (mandatory, 30 min) — program updates, showcases, new tool announcements.\n\n"
            "Quarterly 1:1s per citizen (mandatory, 20 min) — reviews Highlander data, zone compliance, coaching. "
            "Also the moment to flag citizens with low commit activity: a gap between Claude usage and GitHub activity "
            "is the signal to discuss. This is the detection layer — not a wall, but a regular check-in that keeps "
            "everyone accountable without blocking anyone's ability to work.\n\n"
            "Quarterly MCP connector review: review Anthropic's official connector list for additions since last quarter. "
            "Each new connector is evaluated before citizens can use it — managed-mcp.json is updated only after approval. "
            "This is what prevents the approved tool surface from silently expanding over time.\n\n"
            "Ad hoc incident review for any No-Drive Zone violation."
        ),
        "555555", LGRAY,
    ),
]

steps_tbl = doc.add_table(rows=len(steps) + 1, cols=4)
steps_tbl.style = "Table Grid"
col_widths(steps_tbl, [0.38, 1.55, 2.0, 3.84])

for cell, txt in zip(steps_tbl.rows[0].cells, ["#", "Owner", "Action", "Detail"]):
    set_cell_bg(cell, "0F1117")
    cell_write(cell, txt, bold=True, color=WHITE)

for i, (num, owner, title_txt, detail, bg_hex, txt_color) in enumerate(steps, 1):
    row  = steps_tbl.rows[i]
    fill = "F0F4FF" if i % 2 == 1 else "FFFFFF"
    for cell in row.cells:
        set_cell_bg(cell, fill)
    cell_write(row.cells[0], num,       bold=True, color=BLUE)
    cell_write(row.cells[1], owner,     bold=True, color=txt_color)
    cell_write(row.cells[2], title_txt, bold=True, color=DARK)
    # Detail cell: write as multi-line with hyperlinks where needed
    dc = row.cells[3]
    dc.text = ""

    def cell_para(cell, first_flag):
        if first_flag:
            return cell.paragraphs[0]
        return cell.add_paragraph()

    def style_run(p2, text):
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(text)
        r2.font.name  = "Calibri"
        r2.font.size  = Pt(9.5)
        r2.font.color.rgb = LGRAY

    if num == "1":
        lines = [
            ("GitHub Teams is the version control and audit trail infrastructure for the program. "
             "Every citizen's AI-assisted work lives here — feeds Highlander and provides the SOC2 compensating control.",
             False, None),
            ("", False, None),
            ("GitHub Teams:  $4.00/user/month (monthly)  ·  $3.67/user/month (annual).", False, None),
            ("Pricing source: github.com/pricing", True, "https://github.com/pricing"),
            ("", False, None),
            ("Claude Teams (contracted):  $25/seat/month (standard)  ·  $125/seat/month (premium, incl. Claude Code).",
             False, None),
            ("Pricing source: claude.com/pricing", True, "https://claude.com/pricing"),
        ]
        first = True
        for line_txt, is_link, url in lines:
            p2 = cell_para(dc, first); first = False
            p2.paragraph_format.space_after = Pt(1)
            if is_link and url:
                hyperlink_run(p2, line_txt, url, size=9.5, color=(0x2D, 0x6F, 0xF7))
            else:
                style_run(p2, line_txt)
    else:
        first = True
        for line in detail.split("\n"):
            p2 = cell_para(dc, first); first = False
            style_run(p2, line)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — GUARDRAILS
# ══════════════════════════════════════════════════════════════════════════════
h1("Recommended Guardrails")

body(
    "All controls below are confirmed feasible on the Claude Teams plan — no Enterprise upgrade required. "
    "There are two deployment paths. "
    "The admin console (Admin Settings > Claude Code > Managed Settings) is the simpler path: "
    "configure JSON once and it pushes to every authenticated user. This covers permissions, hooks, and sandbox. "
    "It does not cover MCP server control — that requires the file-based path instead. "
    "The file-based path means IT (or an onboarding script) drops a file in a system directory on each machine; "
    "both Claude Code in the terminal and Claude Desktop read from the same file, so one deployment covers both. "
    "CLAUDE.md does not go through either path — it lives in the repo template and loads automatically."
)

note(
    "Source: code.claude.com/docs/en/server-managed-settings  ·  "
    "code.claude.com/docs/en/permissions  ·  code.claude.com/docs/en/sandboxing"
)

g_rows = [
    (
        "managed-settings.json\n\nDeploy via:\nAdmin console\n(recommended)\nor file drop",
        "Sets permission rules for Claude across the whole org: block dangerous commands "
        "(rm -rf, push to main), deny access to sensitive file paths (.env, .ssh/), "
        "restrict outbound network to approved domains. "
        "Also contains hook definitions and sandbox config — all in one file.\n\n"
        "Admin console path: Admin Settings > Claude Code > Managed Settings in claude.ai. "
        "Configure as JSON, applies to all authenticated users automatically. "
        "Available on Teams plan (v2.1.38+, public beta).\n\n"
        "File path (stable alternative): IT drops the file at "
        "/Library/Application Support/ClaudeCode/managed-settings.json (macOS) "
        "or /etc/claude-code/managed-settings.json (Linux). "
        "Applies to both Claude Code in the terminal and Claude Desktop — same file, both products.",
        "Cannot be overridden by the user or by a prompt. "
        "Setting allowManagedPermissionRulesOnly means citizens cannot write their own rules to undo it.",
        "✓ Teams\n\nAdmin console\nor file-based",
    ),
    (
        "managed-mcp.json\n\nDeploy via:\nFile drop only\n(onboarding script)",
        "Controls which MCP servers citizens are allowed to connect. "
        "Set to deny-by-default: only the 26 official Claude connectors, Highlander, and Playwright are permitted. "
        "Citizens cannot add other servers without written approval.\n\n"
        "Cannot be pushed through the admin console — this is a confirmed limitation of the Teams plan. "
        "Must be deployed as a file alongside managed-settings.json. "
        "For 25 people, the practical approach is a one-time onboarding script: "
        "citizen runs it once during setup, it pulls the file from a private GitHub repo "
        "and places it in the correct system directory for their OS.",
        "MCP servers can do a lot — including read files, call APIs, and talk to live systems. "
        "Without this control, a citizen could connect an unapproved server and inadvertently expose data. "
        "Auditors can inspect the allowlist as a documented, enforced policy.",
        "✓ Teams\n\nFile-based only\n(not admin console)",
    ),
    (
        "CLAUDE.md\n\nDeploy via:\nRepo template\n(auto-loads)",
        "A text file committed to every citizen repo template at .claude/CLAUDE.md. "
        "When a citizen opens a project folder in Claude Code or Claude Desktop, "
        "this file loads automatically — no setup required.\n\n"
        "Contents: Driving Zones reminder, no-credential rule, pointer to the AUP, "
        "and the Jensen Huang framing (job = purpose + tasks; AI handles tasks). "
        "Kept short — under ~150 lines — so Claude reliably follows it.\n\n"
        "Note: the admin console can push JSON settings but cannot push .md files. "
        "The repo template approach is the correct mechanism on Teams plan.",
        "Primes every session with the right context before any task begins. "
        "Zero setup burden on citizens — they just clone the repo and it's there.",
        "✓ Teams\n\nRepo template\n(no deployment needed)",
    ),
    (
        "Hooks\n(PreToolUse)\n\nDeploy via:\nAdmin console\nor file drop",
        "Hooks run automatically before Claude executes any tool. "
        "Configured inside managed-settings.json — same deployment as above.\n\n"
        "What we block: rm -rf, git push directly to main or master, "
        "and command patterns that look like credentials or secrets being pasted. "
        "When a block fires, the citizen sees a plain-English message explaining what to do instead "
        "(e.g. 'use a feature branch').\n\n"
        "Setting allowManagedHooksOnly in managed-settings.json ensures citizens cannot "
        "disable or override the managed hooks.",
        "Hard stop on the most common destructive mistakes — deleting work, pushing to production, "
        "accidentally committing a password. Normal work is unaffected.",
        "✓ Teams\n\nPart of\nmanaged-settings.json",
    ),
    (
        "Sandbox\n(Seatbelt / bwrap)\n\nDeploy via:\nAdmin console\nor file drop",
        "OS-level isolation that restricts what Claude can touch. "
        "Filesystem: writes are limited to the citizen's active project folder. "
        "Network: outbound requests go through a proxy that only allows approved domains.\n\n"
        "macOS uses Apple's built-in Seatbelt framework. Linux uses bubblewrap (requires package install). "
        "Not on by default — must be enabled. "
        "Enable org-wide by setting sandbox.failIfUnavailable: true in managed-settings.json "
        "so it turns on for all citizens automatically.",
        "The only layer that genuinely cannot be bypassed — not by the user, not by a prompt injection. "
        "Even if a citizen pastes malicious instructions into Claude, the sandbox prevents "
        "data from leaving the machine or files outside the project from being touched.",
        "✓ Teams\n\nPart of\nmanaged-settings.json",
    ),
    (
        "PostToolUse\naudit log\n\n(Phase 2)",
        "Logs every tool use Claude makes — what action, what file, what command, when. "
        "Written to a local file or posted to an endpoint. Tabled for initial rollout.\n\n"
        "Also configured inside managed-settings.json when ready. "
        "Combined with the GitHub org audit log (which already captures every commit and PR), "
        "this closes the loop on the SOC2 monitoring requirement.",
        "GitHub covers everything after the commit. This covers what happened inside the Claude session. "
        "Together they make a complete audit trail for SOC2 auditors.",
        "✓ Teams\n\n(Phase 2 —\nnot needed\nfor launch)",
    ),
]

g_tbl = doc.add_table(rows=len(g_rows) + 1, cols=4)
g_tbl.style = "Table Grid"
col_widths(g_tbl, [1.65, 2.75, 2.5, 0.87])

for cell, txt in zip(g_tbl.rows[0].cells, ["Control", "What It Does", "Why It Matters", "Teams Plan?"]):
    set_cell_bg(cell, "0F1117")
    cell_write(cell, txt, bold=True, color=WHITE)

for i, (ctrl, what, why, status) in enumerate(g_rows, 1):
    row  = g_tbl.rows[i]
    is_p2 = (i == len(g_rows))
    fill  = "FFF8E1" if is_p2 else ("F0F4FF" if i % 2 == 1 else "FFFFFF")
    for cell in row.cells:
        set_cell_bg(cell, fill)
    ctrl_clr  = GOLD if is_p2 else DARK
    stat_clr  = GOLD if is_p2 else GREEN
    cell_write(row.cells[0], ctrl,   bold=True,  color=ctrl_clr)
    cell_write(row.cells[1], what,   italic=is_p2)
    cell_write(row.cells[2], why,    italic=is_p2)
    cell_write(row.cells[3], status, bold=True,  color=stat_clr)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — ARCHITECTURE DIAGRAM
# ══════════════════════════════════════════════════════════════════════════════
h1("Citizen Workflow Architecture")

body(
    "The diagram below shows how each citizen's work flows up through the system — "
    "from Claude Desktop or Claude Code, through the guardrail stack, into GitHub, "
    "and into Highlander to generate CIO-level metrics automatically."
)

diagram_buf = build_diagram()
doc.add_picture(diagram_buf, width=Inches(7.77))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

p_cap = doc.add_paragraph(
    "Citizens are at the bottom. Work flows upward through the guardrail stack (amber), "
    "into the GitHub org, through Highlander's analyzers, and up to the CIO dashboard. "
    "No manual reporting — every commit is tracked automatically via the GitHub webhook."
)
p_cap.paragraph_format.space_before = Pt(4)
for run in p_cap.runs:
    run.font.name   = "Calibri"
    run.font.size   = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — OUTSTANDING RISKS
# ══════════════════════════════════════════════════════════════════════════════
h1("Outstanding Risks")

body(
    "The controls in this plan provide a strong baseline. The risks below are real but either require "
    "IT or leadership involvement, represent lower-probability scenarios, or are being tabled pending "
    "further program maturity. Each is documented here so technical leaders can decide whether to address "
    "any item before rollout — or accept it as a known, managed risk."
)

risk_rows = [
    (
        "Guardrails don't apply to Claude.ai in a browser or on mobile",
        "All technical controls (managed-settings, MCP allowlist, hooks) are local-client only. "
        "A citizen doing Zone 3 work can simply open claude.ai in a browser tab and bypass everything.",
        "AUP clause: 'All company work must use approved Claude Desktop or Claude Code clients only.' "
        "SSO enforcement via IdP can block unmanaged login paths. MDM web filtering can block claude.ai entirely.",
    ),
    (
        "GitHub audit log retention is only 180 days",
        "SOC2 Type II audits cover a 12-month period. Evidence older than 6 months cannot be produced from GitHub alone.",
        "Scheduled Lambda or GitHub Action exports audit events to S3 monthly. ~4 hrs engineering, negligible storage cost. "
        "Do this before the program is 6 months old.",
    ),
    (
        "SSO federation and MFA for GitHub not confirmed as active",
        "Without SSO, citizen account deprovisioning on departure is fully manual. "
        "Without MFA, a compromised GitHub account has write access to all citizen repos.",
        "IT to confirm SSO is federated with the GitHub org. "
        "Enable MFA requirement at the org level — single checkbox in GitHub org security settings.",
    ),
    (
        "managed-mcp.json can be modified by a local admin",
        "A citizen with local admin rights on their machine can delete or overwrite the MCP lockdown file "
        "and connect unapproved servers.",
        "IT to confirm local admin rights policy for citizen machines. "
        "MDM (already deployed) can enforce file presence and integrity in system directories.",
    ),
    (
        "No secret scanning on citizen repos",
        "A citizen could accidentally hardcode an API key or credential in their code and commit it. "
        "Without secret scanning, this goes undetected until the key is abused.",
        "Enable GitHub secret push protection at the org level. "
        "Included in Teams plan. Add a .gitignore template to the citizen repo template blocking common credential file patterns.",
    ),
    (
        "No formal incident response playbook for No-Drive Zone violations",
        "The first violation will trigger confusion — who investigates, who suspends access, "
        "how is evidence preserved, who communicates with the affected citizen?",
        "Write a one-page IR playbook: detection triggers, access suspension steps, evidence preservation "
        "(GitHub audit export), investigation steps, communication templates, escalation if Evan is unavailable. ~2 hrs.",
    ),
    (
        "Single program lead",
        "Evan handles all MCP approvals, incident response, 1:1s, and quarterly reviews. "
        "One unavailable week stalls the program or leaves a violation without a response path.",
        "Designate a backup program lead informally. Document which decisions only Evan makes vs. which can be delegated.",
    ),
    (
        "Citizens are new to GitHub and may not commit frequently",
        "Highlander's business value metrics depend on commit activity. "
        "Long gaps mean Highlander sees nothing, OKR attribution breaks, and the ROI story is harder to prove.",
        "Highlander can be configured to flag citizens with no commits in 30+ days. "
        "Address in quarterly 1:1s. Course GitHub module must set explicit commit frequency expectations.",
    ),
    (
        "METADATA.yaml completeness is unenforced",
        "Citizens create repos from the template but may leave placeholder OKR text. "
        "Highlander's OKR attribution breaks silently — Highlander can't report what it can't map.",
        "GitHub Actions workflow checks METADATA.yaml on every PR to main — fails if the OKR field is still a placeholder. ~1-2 hrs.",
    ),
    (
        "Data retention policy for Claude sessions is undocumented",
        "Auditors may ask how long Claude retains conversation data. "
        "Without a documented answer, there is a gap in ISO A.8.10 (information deletion).",
        "Pull exact retention terms from the Anthropic Teams contract. "
        "Add a one-paragraph summary to the citizen policies document.",
    ),
]

risk_tbl = doc.add_table(rows=len(risk_rows) + 1, cols=3)
risk_tbl.style = "Table Grid"
col_widths(risk_tbl, [2.2, 2.5, 3.07])

for cell, txt in zip(risk_tbl.rows[0].cells, ["The Gap", "Why It Matters", "What Could Be Done"]):
    set_cell_bg(cell, "0F1117")
    cell_write(cell, txt, bold=True, color=WHITE)

for i, (gap, why, what) in enumerate(risk_rows, 1):
    row  = risk_tbl.rows[i]
    fill = "FFF8E1" if i % 2 == 1 else "FFFDF0"
    for cell in row.cells:
        set_cell_bg(cell, fill)
    cell_write(row.cells[0], gap,  bold=True,  color=RGBColor(0xB8, 0x76, 0x00))
    cell_write(row.cells[1], why)
    cell_write(row.cells[2], what)

divider()

# ─── Footer ───────────────────────────────────────────────────────────────────
p_foot = doc.add_paragraph(
    "Full documentation: cio-plan-citizen-ai-engineer.md  ·  "
    "Master deck: Citizen AI Engineer Program — Master Deck V1.pptx  ·  "
    "Course: ialta.sharepoint.com/sites/CitizenAI"
)
p_foot.paragraph_format.space_before = Pt(6)
for run in p_foot.runs:
    run.font.name      = "Calibri"
    run.font.size      = Pt(8.5)
    run.font.italic    = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save(OUT)
print(f"✓  Saved: {OUT}")
