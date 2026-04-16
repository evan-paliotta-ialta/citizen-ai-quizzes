#!/usr/bin/env python3
"""Generate: Citizen AI Engineer Program — Policies, Procedures & Agreements.docx"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
OUT  = os.path.join(BASE, "Citizen AI Engineer Program — Policies, Procedures & Agreements.docx")

# ─── Colours ──────────────────────────────────────────────────────────────────
DARK  = RGBColor(0x0F, 0x11, 0x17)
BLUE  = RGBColor(0x2D, 0x6F, 0xF7)
HBLU  = RGBColor(0x1A, 0x48, 0xBB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY = RGBColor(0x44, 0x44, 0x44)
GREEN = RGBColor(0x1A, 0x7A, 0x3E)
GOLD  = RGBColor(0xB8, 0x76, 0x00)
RED   = RGBColor(0xC0, 0x39, 0x2B)
MGRAY = RGBColor(0x77, 0x77, 0x77)

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT SETUP
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after  = Pt(3)
style.paragraph_format.space_before = Pt(0)


# ─── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_str):
    tcPr = cell._element.get_or_add_tcPr()
    shd  = tcPr.makeelement(qn("w:shd"),
                             {qn("w:fill"): hex_str, qn("w:val"): "clear"})
    tcPr.append(shd)


def cell_write(cell, text, size=9.5, bold=False, color=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.name   = "Calibri"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color or LGRAY


def col_widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.font.name  = "Calibri"
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = BLUE
    return p


def h2(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(10.5)
    run.font.bold  = True
    run.font.color.rgb = color or DARK
    return p


def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(10)
    run.font.color.rgb = LGRAY
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.font.name = "Calibri"
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = DARK
        r2 = p.add_run(text)
        r2.font.name = "Calibri"
        r2.font.size = Pt(10)
        r2.font.color.rgb = LGRAY
    else:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = LGRAY
    return p


def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.2)
    run = p.add_run(text)
    run.font.name   = "Calibri"
    run.font.size   = Pt(9)
    run.font.italic = True
    run.font.color.rgb = MGRAY
    return p


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 100)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def checkbox_item(text, bold_prefix=None):
    """Render a checklist item with a checkbox character."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.2)
    box_run = p.add_run("☐  ")
    box_run.font.name = "Calibri"
    box_run.font.size = Pt(11)
    box_run.font.color.rgb = LGRAY
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.font.name = "Calibri"
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = DARK
        r2 = p.add_run(text)
        r2.font.name = "Calibri"
        r2.font.size = Pt(10)
        r2.font.color.rgb = LGRAY
    else:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = LGRAY
    return p


def numbered(num, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.25)
    r0 = p.add_run(f"{num}.  ")
    r0.font.name  = "Calibri"
    r0.font.size  = Pt(10)
    r0.font.bold  = True
    r0.font.color.rgb = BLUE
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.font.name  = "Calibri"
        r1.font.size  = Pt(10)
        r1.font.bold  = True
        r1.font.color.rgb = DARK
        r2 = p.add_run(text)
        r2.font.name  = "Calibri"
        r2.font.size  = Pt(10)
        r2.font.color.rgb = LGRAY
    else:
        run = p.add_run(text)
        run.font.name  = "Calibri"
        run.font.size  = Pt(10)
        run.font.color.rgb = LGRAY
    return p


# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("Citizen AI Engineer Program", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in title.runs:
    run.font.color.rgb = DARK
    run.font.size      = Pt(18)
    run.font.name      = "Calibri"

sub = doc.add_paragraph("Policies, Procedures & Agreements  ·  April 2026  ·  Confidential")
for run in sub.runs:
    run.font.name      = "Calibri"
    run.font.size      = Pt(10)
    run.font.color.rgb = MGRAY
    run.font.italic    = True

divider()

intro_p = doc.add_paragraph()
intro_p.paragraph_format.space_before = Pt(4)
intro_p.paragraph_format.space_after  = Pt(10)
run = intro_p.add_run(
    "The company-wide AI Acceptable Use Policy (AUP) applies to every employee. "
    "This document is the additional layer that applies specifically to you as a Citizen AI Engineer. "
    "It translates the AUP into plain English, adds the program-specific procedures and expectations, "
    "and includes the agreements you sign when you join the program. "
    "Read it carefully. You sign it after completing the certification course."
)
run.font.name      = "Calibri"
run.font.size      = Pt(10.5)
run.font.color.rgb = LGRAY

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — THE AUP IN PLAIN ENGLISH
# ══════════════════════════════════════════════════════════════════════════════
h1("1. The Company AI Policy — What It Means for You")

body(
    "The full AI Acceptable Use Policy is a formal legal document. Here are the rules that matter most "
    "in your day-to-day work as a Citizen AI Engineer, stated plainly:"
)

aup_rules = [
    ("The data you're working with determines the rules.",
     "Know your zone before you start. The Driving Zones reference in Section 3 is the decision tool."),
    ("Claude Teams is the only approved AI tool for company work.",
     "No personal Claude.ai accounts. No ChatGPT. No other AI tools for any work involving company data. "
     "The approved client is Claude Desktop or Claude Code — not Claude.ai in a browser or on your phone."),
    ("Never paste credentials, API keys, or secrets into Claude.",
     "This includes database passwords, tokens, private keys, and anything that authenticates a system. "
     "Ever. Under any circumstances. Regardless of zone."),
    ("AI output is a draft — you are responsible for what you act on or share.",
     "Treat everything Claude generates as a starting point, not a finished product. "
     "Fact-check anything you'll publish, distribute, or act on."),
    ("Client data and confidential information require Zone 3 process.",
     "See Section 3. Additional steps apply. Do not skip them."),
    ("Report any incident immediately.",
     "Accidental data paste, unexpected AI behavior, suspicious output, anything that feels wrong — "
     "contact IT Security (and Evan) the same business day. Do not delete prompts or outputs first."),
    ("Violations are graduated but serious.",
     "First offense: retraining. Repeated or serious violations: access revoked. "
     "Severe violations (intentional data exfiltration, deceptive content): termination and possible legal action. "
     "The company takes this seriously. So should you."),
]

for bold_text, detail in aup_rules:
    bullet(detail, bold_prefix=bold_text)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — RISK TIERS & DRIVING ZONES
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Risk Tiers and the Driving Zones")

body(
    "The AUP uses a four-tier risk framework (Tier 1–4). The Citizen AI Engineer program maps these tiers "
    "to the Driving Zones model. Both systems say the same thing — the data you're working with determines "
    "the process required. Here is how they align:"
)

zone_tbl = doc.add_table(rows=6, cols=4)
zone_tbl.style = "Table Grid"
col_widths(zone_tbl, [1.7, 1.3, 1.7, 3.07])

zone_headers = ["Driving Zone", "AUP Tier", "Data Type", "What's Required"]
for j, hdr in enumerate(zone_headers):
    set_cell_bg(zone_tbl.rows[0].cells[j], "0F1117")
    cell_write(zone_tbl.rows[0].cells[j], hdr, bold=True, color=WHITE)

zone_data = [
    ("Zone 1 — Private Property",  "Tier 1 — Minimal",
     "Public data only. No company info, no client names, no internal terminology.",
     "No governance required. Push to GitHub anyway — Highlander tracks all citizen activity."),
    ("Zone 2 — Residential Streets", "Tier 2 — Low",
     "Internal company data (non-confidential). No PII. No client data.",
     "GitHub repo required. Branch protection on. No direct push to main. Code review before merge."),
    ("Zone 3 — The Highway", "Tier 3 — Elevated",
     "Confidential data, PII, client data, or any live system accessed via MCP.",
     "Everything in Zone 2, plus: approved MCP servers only, full commit documentation, manager review."),
    ("Zone 3 (Escalated) / No-Drive", "Tier 4 — High Risk",
     "Agentic workflows with autonomous decisions, or anything touching legal/financial/employment outcomes.",
     "Executive-level approval required before deployment. Likely outside citizen scope — escalate to Evan."),
    ("No-Drive Zone", "Prohibited (AUP §7)",
     "Credentials, secrets, unapproved tools, personal AI accounts for work, bypassing guardrails.",
     "Strictly prohibited. Violation triggers immediate license suspension and review."),
]

zone_fills = ["F0F4FF", "FFFFFF", "FFF3CD", "FFE8E8", "FFD6D6"]
zone_text_colors = [DARK, DARK, DARK, RED, RED]

for i, (zone, tier, data, req) in enumerate(zone_data, 1):
    row = zone_tbl.rows[i]
    for cell in row.cells:
        set_cell_bg(cell, zone_fills[i - 1])
    cell_write(row.cells[0], zone, bold=True, color=zone_text_colors[i - 1])
    cell_write(row.cells[1], tier, bold=True, color=MGRAY)
    cell_write(row.cells[2], data)
    cell_write(row.cells[3], req)

doc.add_paragraph()
note(
    "Decision rule: The data determines the zone. Any MCP server connecting to a live system is "
    "automatically Zone 3, regardless of what data you think is involved. When in doubt, go up a zone."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — CITIZEN-SPECIFIC OBLIGATIONS
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Citizen AI Engineer — Specific Obligations")

body(
    "The following apply to program participants only. These are in addition to the company-wide AUP."
)

h2("All AI-assisted work goes in GitHub")
bullet("No local-only work. Every project, every output, every experiment — commit it to a GitHub repo.")
bullet("Zone 1 work included. Even if you're just learning, a repo entry creates a record Highlander can use.")
bullet("Push frequently. At minimum weekly when actively working. This is how the program proves its value.")

h2("GitHub account requirements")
bullet("Your GitHub account must be registered with your company email address (@verivend.com / @ialta.com).")
bullet("MFA (multi-factor authentication) must be enabled on your GitHub account before you receive access.")
bullet("Repo naming convention: citizen-<team>-<project> (e.g., citizen-ops-arr-report).")
bullet("Topic tag: every repo must be tagged citizen-ai so Highlander can segment it.")

h2("METADATA.yaml in every repo")
body(
    "Every citizen repo must contain a completed METADATA.yaml file. This is how Highlander links your "
    "work to an OKR. A repo without a completed METADATA.yaml is invisible to the program's tracking.",
    indent=True
)
note(
    "Required fields: owner (your company email), team, okr (the OKR this work addresses), "
    "business_problem (one sentence), zone (1, 2, or 3)."
)

h2("Approved clients only")
bullet(
    "Use Claude Desktop or Claude Code only for company work. "
    "Claude.ai in your browser and the Claude mobile app do not have the program's guardrails — "
    "the managed settings, the MCP allowlist, and the hooks only apply to the approved desktop clients."
)

h2("MCP server policy")
bullet("Only pre-approved MCP servers may be connected. The approved list is in Section 4.")
bullet("Playwright MCP is approved — but any session using Playwright is automatically Zone 3.")
bullet("Any MCP server not on the approved list requires written approval from the program lead before first use.")
bullet("You may not add unapproved MCP servers. This is a No-Drive Zone violation.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — APPROVED MCP SERVERS
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Approved MCP Servers")

body(
    "The following MCP servers are pre-approved for citizen use. The list is maintained by the program lead "
    "and reviewed quarterly. Any server not on this list requires written approval before first use."
)

mcp_tbl = doc.add_table(rows=4, cols=3)
mcp_tbl.style = "Table Grid"
col_widths(mcp_tbl, [2.3, 2.0, 3.47])

for j, hdr in enumerate(["Server", "Type", "Notes"]):
    set_cell_bg(mcp_tbl.rows[0].cells[j], "0F1117")
    cell_write(mcp_tbl.rows[0].cells[j], hdr, bold=True, color=WHITE)

mcp_data = [
    ("26 official Claude Desktop native connectors",
     "Anthropic-official only",
     "The connectors listed in the Claude Desktop native connector list. "
     "Community-created servers are NOT included even if they appear similar. "
     "The approved list is reviewed quarterly — new additions require evaluation before use."),
    ("Highlander MCP",
     "Internal (iAltA)",
     "The company's internal analytics platform. Any session using Highlander is Zone 3 — "
     "it connects to live company data."),
    ("Playwright MCP",
     "Official (Microsoft)",
     "Browser automation. Approved, but ANY session using Playwright is automatically Zone 3. "
     "Be aware: malicious content on web pages you visit can attempt to inject instructions into Claude's context."),
]

for i, (srv, typ, notes) in enumerate(mcp_data, 1):
    row = mcp_tbl.rows[i]
    fill = "F0F4FF" if i % 2 == 1 else "FFFFFF"
    for cell in row.cells:
        set_cell_bg(cell, fill)
    cell_write(row.cells[0], srv,   bold=True, color=DARK)
    cell_write(row.cells[1], typ,   bold=True, color=BLUE)
    cell_write(row.cells[2], notes)

doc.add_paragraph()
note(
    "All others require written approval from the program lead (Evan) before connecting. "
    "Connecting an unapproved MCP server is a No-Drive Zone violation."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — OFFBOARDING CHECKLIST
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Offboarding Checklist")

body(
    "This checklist is executed by IT and the program lead when a Citizen AI Engineer leaves the program — "
    "whether through departure, role change, license revocation, or voluntary exit. "
    "All steps must be completed within one business day of the offboarding trigger."
)

h2("Access Removal", color=RED)
checkbox_item("Remove from Claude Teams org (Admin Settings > Members)")
checkbox_item("Remove from GitHub org / citizen team")
checkbox_item("Confirm no active Claude API keys or tokens associated with the account")

h2("Repo Continuity", color=HBLU)
checkbox_item("Transfer ownership of all citizen repos to the program lead or a designated successor")
checkbox_item("Resolve or hand off all open pull requests")
checkbox_item("Archive repos if no active successor will continue the work")
checkbox_item("Confirm METADATA.yaml owner field is updated to the new owner's email")

h2("Tracking & Certification", color=GREEN)
checkbox_item("Remove from Highlander engineer_email_mapping table")
checkbox_item("Mark certification as inactive in the Microsoft Forms completion tracker (Excel)")
checkbox_item("Note departure date in the program participation log")

h2("Confirmation", color=MGRAY)
checkbox_item("Confirm all steps above are complete", bold_prefix="Sign-off:")
note("Completed by:  ________________   Date:  ________________   Verified by:  ________________")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — INCIDENT RESPONSE
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Incident Response — What to Do If Something Goes Wrong")

body(
    "An AI incident includes: accidentally pasting confidential data or PII into Claude, "
    "discovering that a session produced output containing sensitive information, "
    "unexpected AI behavior (Claude taking actions you didn't request), "
    "or any suspicion that a policy violation occurred. "
    "The AUP (Section 13) governs the full investigation process. Here is what you do immediately:"
)

numbered(1, "Stop what you're doing. Do not continue the session.", bold_prefix="Stop.")
numbered(2, "Do not delete anything. Do not close the conversation, delete prompts, or clear outputs. "
         "Evidence must be preserved exactly as it is.", bold_prefix="Preserve.")
numbered(3, "Contact IT Security and the program lead (Evan) within the same business day. "
         "Do not wait to see if it becomes a problem.", bold_prefix="Report.")
numbered(4, "IT Security leads the investigation. Your job is to cooperate fully and make yourself available.", bold_prefix="Cooperate.")
numbered(5, "No retaliation for reporting in good faith. You will not be penalized for surfacing a concern promptly.", bold_prefix="Safe to report.")

doc.add_paragraph()
body(
    "Violations are handled per the AUP graduated framework: minor violations result in retraining; "
    "repeated or careless violations result in temporary access restriction; serious violations result in "
    "full access suspension; severe or intentional violations may result in termination and legal action. "
    "A No-Drive Zone violation triggers immediate license suspension pending investigation."
)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — ACKNOWLEDGMENT & SIGNATURE
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Acknowledgment & Signature")

body(
    "By signing below, you confirm that you have read and understood this document in full, "
    "and that you agree to the obligations and procedures it contains."
)

sig_tbl = doc.add_table(rows=9, cols=2)
sig_tbl.style = "Table Grid"
col_widths(sig_tbl, [3.9, 3.87])

sig_checks = [
    "I have read and understood the company-wide AI Acceptable Use Policy (v2.0).",
    "I have completed the Citizen AI Engineer certification course and passed the final exam.",
    "I understand the Driving Zones model and will correctly classify my work before starting any session.",
    "I understand that all company AI work must be done in Claude Desktop or Claude Code — not Claude.ai web or mobile.",
    "I will commit all AI-assisted work to my GitHub repo. No local-only work.",
    "I will never input credentials, API keys, secrets, or passwords into Claude under any circumstances.",
    "I understand that my commit activity is tracked by Highlander and attributed to my OKRs.",
    "I understand that a No-Drive Zone violation results in immediate license suspension pending review, "
    "and may result in permanent revocation.",
]

for i, check_text in enumerate(sig_checks):
    row = sig_tbl.rows[i]
    fill = "F0F4FF" if i % 2 == 0 else "FFFFFF"
    set_cell_bg(row.cells[0], fill)
    set_cell_bg(row.cells[1], fill)
    cb_p = row.cells[0].paragraphs[0]
    cb_p.paragraph_format.space_before = Pt(3)
    cb_p.paragraph_format.space_after  = Pt(3)
    r_box = cb_p.add_run("☐  ")
    r_box.font.name = "Calibri"
    r_box.font.size = Pt(11)
    r_box.font.color.rgb = LGRAY
    r_txt = cb_p.add_run(check_text)
    r_txt.font.name  = "Calibri"
    r_txt.font.size  = Pt(9.5)
    r_txt.font.color.rgb = LGRAY
    # Right cell: blank (acknowledgment column)
    row.cells[1].text = ""

# Last row — signature fields
last_row = sig_tbl.rows[8]
set_cell_bg(last_row.cells[0], "EEF2FF")
set_cell_bg(last_row.cells[1], "EEF2FF")
cell_write(last_row.cells[0], "Full Name:  _______________________________\n\nTeam / Department:  _______________________________", bold=False)
cell_write(last_row.cells[1], "Signature:  _______________________________\n\nDate:  _______________________________", bold=False)

doc.add_paragraph()
note(
    "Return this signed document to the program lead (Evan) upon completion of the certification course. "
    "License access is granted after receipt of this signed agreement and confirmation of course completion."
)

# ─── Footer ───────────────────────────────────────────────────────────────────
p_foot = doc.add_paragraph(
    "Citizen AI Engineer Program  ·  Policies, Procedures & Agreements  ·  "
    "Companion to: Corporate AI Use Policy v2.0  ·  iAltA Private Markets / Verivend Inc."
)
p_foot.paragraph_format.space_before = Pt(8)
for run in p_foot.runs:
    run.font.name      = "Calibri"
    run.font.size      = Pt(8.5)
    run.font.italic    = True
    run.font.color.rgb = MGRAY

doc.save(OUT)
print(f"✓  Saved: {OUT}")
