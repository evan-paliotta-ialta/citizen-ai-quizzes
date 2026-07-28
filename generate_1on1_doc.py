#!/usr/bin/env python3
"""
generate_1on1_doc.py — builds a one-page docx summary of a quarterly 1:1,
for upload to the restricted "Program Admin" SharePoint library.

Run: python3 generate_1on1_doc.py --json '{"name": "...", ...}' --out /tmp/out.docx
"""
import argparse
import json

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x00, 0x0D, 0x2D)
BLUE = RGBColor(0x00, 0x42, 0xE0)


def add_heading(doc, text, size=16, color=NAVY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_field(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(value or "—")
    r2.font.size = Pt(11)
    return p


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--json", required=True, help="JSON blob with the 1:1 fields")
    ap.add_argument("--out", required=True, help="Output .docx path")
    args = ap.parse_args()

    data = json.loads(args.json)

    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(f"1:1 — {data['name']}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"{data['quarter']}  |  {data['date']}")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = BLUE

    doc.add_paragraph()

    add_field(doc, "Email", data.get("email", ""))
    add_field(doc, "Capability Level", data.get("capability", ""))
    add_field(doc, "Prior Goal Status", data.get("prior_status", ""))

    doc.add_paragraph()
    add_heading(doc, "What they built this quarter")
    doc.add_paragraph(data.get("built", "—"))

    add_heading(doc, "Blocker")
    doc.add_paragraph(data.get("blocker", "—"))

    add_heading(doc, "Goal set for next quarter")
    doc.add_paragraph(data.get("goal", "—"))

    if data.get("notes"):
        add_heading(doc, "Notes")
        doc.add_paragraph(data["notes"])

    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Private coaching conversation — not recorded. Restricted to program administrators."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(args.out)
    print(f"saved {args.out}")


if __name__ == "__main__":
    main()
